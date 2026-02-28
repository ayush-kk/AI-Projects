import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings

logger = logging.getLogger(__name__)


def generate_word(data: dict, user_id: int) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    title = data.get("title", "Document")
    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    subtitle = data.get("subtitle")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_paragraph("")

    sections = data.get("sections", [])
    for section in sections:
        section_title = section.get("heading", "")
        if section_title:
            doc.add_heading(section_title, level=1)

        content = section.get("content", "")
        if content:
            doc.add_paragraph(content)

        bullet_points = section.get("bullets", [])
        for bullet in bullet_points:
            doc.add_paragraph(bullet, style="List Bullet")

        table_data = section.get("table")
        if table_data:
            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if headers:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Light Grid Accent 1"
                for j, header in enumerate(headers):
                    cell = table.rows[0].cells[j]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                for i, row in enumerate(rows):
                    for j, value in enumerate(row):
                        table.rows[i + 1].cells[j].text = str(value)
                doc.add_paragraph("")

    output_dir = Path(settings.generated_dir) / str(user_id)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = str(output_dir / filename)
    doc.save(filepath)
    return filepath
